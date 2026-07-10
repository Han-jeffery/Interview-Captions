#!/usr/bin/env python3
import json
import sys
from pathlib import Path


def extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_pdf(path: Path) -> str:
    try:
        import pdfplumber

        parts = []
        with pdfplumber.open(path) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    parts.append(f"[Page {index}]\n{text}")
        if parts:
            return "\n\n".join(parts)
    except Exception:
        pass

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"[Page {index}]\n{text}")
    return "\n\n".join(parts)


def main() -> int:
    if len(sys.argv) != 3:
        print(json.dumps({"ok": False, "error": "Usage: extract_material.py <path> <ext>"}))
        return 2

    path = Path(sys.argv[1])
    ext = sys.argv[2].lower().lstrip(".")

    try:
        if ext == "docx":
            text = extract_docx(path)
        elif ext == "pdf":
            text = extract_pdf(path)
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")

        print(json.dumps({"ok": True, "text": text}, ensure_ascii=False))
        return 0
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
